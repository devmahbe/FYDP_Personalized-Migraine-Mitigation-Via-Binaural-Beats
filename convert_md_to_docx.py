"""
Convert Markdown files to .docx (Word) format, preserving:
- Headings (H1-H4)
- Bold, italic, inline code
- Tables (with borders)
- Code blocks / diagrams (monospaced, shaded background)
- Block quotes
- Bullet/numbered lists
- Horizontal rules
- Math formulas (rendered as text)
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_paragraph_shading(paragraph, color_hex):
    """Set background shading for a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    pPr.append(shading_elm)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_text(paragraph, text, bold=False, italic=False, code=False, size=None, color=None):
    """Add a run of text with formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def parse_inline_formatting(paragraph, text, base_bold=False, base_italic=False, base_size=None):
    """Parse inline markdown formatting: **bold**, *italic*, `code`, $$math$$, $math$."""
    # Pattern to match inline formatting elements
    # Order matters: **bold** before *italic*, $$ before $
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'    # ***bold italic***
        r'|(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'           # *italic*
        r'|(`([^`]+?)`)'          # `code`
        r'|(\$\$(.+?)\$\$)'      # $$math$$
        r'|(\$(.+?)\$)'           # $math$
    )
    
    last_end = 0
    for m in pattern.finditer(text):
        # Add text before the match
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            if plain:
                add_formatted_text(paragraph, plain, bold=base_bold, italic=base_italic, size=base_size)
        
        if m.group(2):  # ***bold italic***
            add_formatted_text(paragraph, m.group(2), bold=True, italic=True, size=base_size)
        elif m.group(4):  # **bold**
            add_formatted_text(paragraph, m.group(4), bold=True, italic=base_italic, size=base_size)
        elif m.group(6):  # *italic*
            add_formatted_text(paragraph, m.group(6), bold=base_bold, italic=True, size=base_size)
        elif m.group(8):  # `code`
            add_formatted_text(paragraph, m.group(8), code=True, size=base_size)
        elif m.group(10):  # $$math$$
            add_formatted_text(paragraph, m.group(10), italic=True, size=base_size,
                             color=RGBColor(0x1A, 0x23, 0x7E))
        elif m.group(12):  # $math$
            add_formatted_text(paragraph, m.group(12), italic=True, size=base_size,
                             color=RGBColor(0x1A, 0x23, 0x7E))
        
        last_end = m.end()
    
    # Add remaining text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            add_formatted_text(paragraph, remaining, bold=base_bold, italic=base_italic, size=base_size)


def parse_table(lines):
    """Parse markdown table lines into list of rows (each row = list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            # Skip separator rows (containing only dashes, colons, spaces)
            if all(re.match(r'^[\-:]+$', c) or c == '' for c in cells):
                continue
            rows.append(cells)
    return rows


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if stripped == '':
            i += 1
            continue
        
        # Horizontal rule
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            # Add a thin horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue
        
        # Code block (``` or ````  fenced)
        code_match = re.match(r'^(`{3,4})(\w*)', stripped)
        if code_match:
            fence = code_match.group(1)
            code_lines = []
            i += 1
            while i < len(lines):
                if lines[i].strip().startswith(fence):
                    i += 1
                    break
                code_lines.append(lines[i])
                i += 1
            
            # Add code block as a single paragraph with monospace font and shading
            code_text = '\n'.join(code_lines)
            if code_text.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Cm(0.5)
                set_paragraph_shading(p, "F5F5F5")
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue
        
        # Math block ($$...$$) spanning multiple lines
        if stripped.startswith('$$') and not stripped.endswith('$$'):
            math_lines = [stripped[2:]]
            i += 1
            while i < len(lines):
                if lines[i].strip().endswith('$$'):
                    math_lines.append(lines[i].strip().rstrip('$'))
                    i += 1
                    break
                math_lines.append(lines[i].strip())
                i += 1
            math_text = ' '.join(math_lines).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(math_text)
            run.font.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
            continue
        
        # Single-line math block $$...$$
        math_single = re.match(r'^\$\$(.+)\$\$$', stripped)
        if math_single:
            math_text = math_single.group(1).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(math_text)
            run.font.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
            i += 1
            continue
        
        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Map markdown levels to Word heading levels
            word_level = min(level, 4)
            h = doc.add_heading(level=word_level)
            # Parse inline formatting in heading
            parse_inline_formatting(h, heading_text, base_bold=True)
            i += 1
            continue
        
        # Table detection
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            
            rows = parse_table(table_lines)
            if rows and len(rows) >= 1:
                num_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                add_table_borders(table)
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            
                            # Header row styling
                            if row_idx == 0:
                                set_cell_shading(cell, "E8EAF6")
                                parse_inline_formatting(p, cell_text, base_bold=True, base_size=10)
                            else:
                                parse_inline_formatting(p, cell_text, base_size=10)
                
                # Add spacing after table
                doc.add_paragraph()
            continue
        
        # Block quote
        if stripped.startswith('>'):
            quote_text = stripped.lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add left border via paragraph borders
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="12" w:space="4" w:color="4A90D9"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            parse_inline_formatting(p, quote_text, base_italic=True)
            i += 1
            continue
        
        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            text = num_match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, text)
            i += 1
            continue
        
        # Bullet list (- or *)
        bullet_match = re.match(r'^[-*]\s+(.+)', stripped)
        if bullet_match:
            text = bullet_match.group(1).strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, text)
            i += 1
            continue
        
        # Bold line by itself (like **text**)
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        parse_inline_formatting(p, stripped)
        i += 1
    
    # Save
    doc.save(docx_path)
    print(f"  Saved: {docx_path}")


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    docs_dir = os.path.join(base_dir, 'docs')
    
    files_to_convert = [
        ('PREPROCESSING_EASY_GUIDE_AND_DEFENSE_QA.md', 'PREPROCESSING_EASY_GUIDE_AND_DEFENSE_QA.docx'),
        ('RESEARCH_PAPER_DOCUMENTATION.md', 'RESEARCH_PAPER_DOCUMENTATION.docx'),
    ]
    
    for md_name, docx_name in files_to_convert:
        md_path = os.path.join(docs_dir, md_name)
        docx_path = os.path.join(docs_dir, docx_name)
        
        if not os.path.exists(md_path):
            print(f"  ERROR: {md_path} not found!")
            continue
        
        print(f"Converting: {md_name}")
        convert_md_to_docx(md_path, docx_path)
    
    print("\nDone! Both .docx files have been generated in the docs/ folder.")


if __name__ == '__main__':
    main()
